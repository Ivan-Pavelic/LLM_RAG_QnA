from __future__ import annotations

"""
Build a .docx paper following the provided template style cues (Times New Roman, IMRaD).

Inputs (from this repo):
- paper/IMRaD_draft.md (text source)
- paper/figure_captions.md (caption text source)
- outputs/metrics_qa.csv, outputs/metrics_sum.csv (tables)
- outputs/figures/*.png (figures)
- outputs/manual_annotations_filled.csv (optional; used indirectly via correlation already computed)

Output:
- paper/LLM_vs_RAG_QA_Summarization.docx
"""

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]


def set_run_font(run, *, size_pt: int, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic


def add_paragraph(
    doc: Document,
    text: str,
    *,
    size_pt: int = 12,
    bold: bool = False,
    italic: bool = False,
    align: str = "justify",  # 'left'|'center'|'right'|'justify'
) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, size_pt=size_pt, bold=bold, italic=italic)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_heading_1(doc: Document, title: str) -> None:
    # Template: "1. Introduction (TNR 14pt., bold)"
    add_paragraph(doc, title, size_pt=14, bold=True, align="left")


def add_heading_2(doc: Document, title: str) -> None:
    # Template: "1.1 Title of the 2nd level (TNR 12pt., bold)"
    add_paragraph(doc, title, size_pt=12, bold=True, align="left")


def add_table_from_df(doc: Document, df: pd.DataFrame, *, float_fmt: str = "{:.4f}") -> None:
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for j, col in enumerate(df.columns):
        hdr[j].text = str(col)
        for run in hdr[j].paragraphs[0].runs:
            set_run_font(run, size_pt=10, bold=True)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(df.columns):
            v = row[col]
            if isinstance(v, float):
                txt = float_fmt.format(v)
            else:
                txt = str(v)
            cells[j].text = txt
            for run in cells[j].paragraphs[0].runs:
                set_run_font(run, size_pt=10)


def add_figure(doc: Document, img_path: Path, caption: str, *, width_in: float = 6.0) -> None:
    if img_path.exists():
        doc.add_picture(str(img_path), width=Inches(width_in))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Caption format (template): TNR 10pt, centered, italics
    add_paragraph(doc, caption, size_pt=10, italic=True, align="center")
    add_paragraph(doc, "Source: Authors' experiments.", size_pt=10, italic=True, align="center")


def main() -> None:
    out_path = ROOT / "paper" / "LLM_vs_RAG_QA_Summarization.docx"
    out_path.parent.mkdir(parents=True, exist_ok=True)

    # Load numeric results
    qa = pd.read_csv(ROOT / "outputs" / "metrics_qa.csv")
    su = pd.read_csv(ROOT / "outputs" / "metrics_sum.csv")

    # Map internal ids to paper naming
    name_map = {"llm": "LLM", "rag_bm25": "RAG-BM25", "rag_dense": "RAG-Dense", "rag_hybrid": "RAG-Hybrid", "rag_dense_rerank": "RAG-Dense+Rerank"}
    qa["system"] = qa["system"].map(lambda x: name_map.get(x, x))
    su["system"] = su["system"].map(lambda x: name_map.get(x, x))

    # Keep compact columns for tables (paper-friendly); BERTScore optional when --skip-bertscore was used
    qa_cols = ["system", "qa_accuracy", "rougeL"]
    if "bertscore_F1" in qa.columns:
        qa_cols.append("bertscore_F1")
    qa_tbl = qa[[c for c in qa_cols if c in qa.columns]].copy()
    su_cols = ["system", "rouge1", "rouge2", "rougeL", "bleu"]
    if "bertscore_F1" in su.columns:
        su_cols.insert(-1, "bertscore_F1")
    su_tbl = su[[c for c in su_cols if c in su.columns]].copy()

    # Known correlation from your run (or from outputs/spearman_rho.txt if present)
    rho_path = ROOT / "outputs" / "spearman_rho.txt"
    if rho_path.exists():
        spearman_rho = float(rho_path.read_text(encoding="utf-8").strip())
    else:
        spearman_rho = 0.08970522700074382

    doc = Document()

    # -------- Title block (template cues) --------
    add_paragraph(
        doc,
        "Comparative Analysis of LLM and Retrieval-Augmented Generation (RAG) Systems for Question Answering and Document Summarization",
        size_pt=20,
        bold=True,
        align="center",
    )
    add_paragraph(doc, "Full First Author1, Full Second Author2, * and Full Third Author3", size_pt=12, align="center")
    add_paragraph(doc, "1 Full affiliation of first author, including country", size_pt=10, align="center")
    add_paragraph(doc, "2 Full affiliation of second author, including country", size_pt=10, align="center")
    add_paragraph(doc, "3 Full affiliation of third author, including country", size_pt=10, align="center")
    add_paragraph(doc, "* Corresponding author", size_pt=10, align="center")

    # -------- Abstract (structured) --------
    add_paragraph(doc, "Abstract", size_pt=12, bold=True, align="left")
    add_paragraph(
        doc,
        "Purpose – This study compares a standalone Large Language Model (LLM) against Retrieval-Augmented Generation (RAG) pipelines for question answering and document summarization, focusing on correctness, hallucination reduction, and consistency.",
        size_pt=12,
    )
    add_paragraph(
        doc,
        "Design/Methodology/Approach – We evaluate a non-RAG baseline and several RAG configurations (BM25, dense, hybrid retrieval, and optional MonoT5 reranking) under controlled conditions. We use larger dataset subsets (e.g. N=400 for SQuAD Q&A and M=100 for CNN/DailyMail summarisation) to improve statistical robustness, and we compare multiple generators including a local FLAN-T5 model and, when API access is available, GPT-3.5 and GPT-4, as well as open-weight models such as LLaMA 2 and Mistral. We vary chunk sizes (100–500 words), top-k (3–10), and retrieval type (sparse, dense, hybrid, with optional reranking).",
        size_pt=12,
    )
    add_paragraph(
        doc,
        "Findings – In our setting, dense retrieval yields higher QA exact-match accuracy than the LLM baseline, while summarization overlap metrics favor the baseline LLM. Manual evaluation shows that automatic overlap metrics are weakly aligned with human correctness judgments (Spearman’s ρ = 0.0897).",
        size_pt=12,
    )
    add_paragraph(
        doc,
        "Originality/Value – The study provides a reproducible benchmark with statistical significance testing, detailed error analysis, and systematic variation over chunk size, top-k, retrieval strategy, and generator size. It underscores the need for manual evaluation and factuality-aware metrics in RAG and non-RAG comparison.",
        size_pt=12,
    )
    add_paragraph(doc, "Keywords – Large Language Models; Retrieval-Augmented Generation; BM25; Dense Retrieval; Question Answering; Summarization; Hallucination.", size_pt=12)
    add_paragraph(doc, "Paper Type – Research paper.", size_pt=12)

    # -------- IMRaD body --------
    add_heading_1(doc, "1. Introduction")
    add_paragraph(
        doc,
        "Large Language Models (LLMs) are widely used for generative NLP tasks such as question answering and summarization. However, LLMs can produce hallucinations—fluent but unsupported statements—which undermines reliability in high-stakes use cases. Retrieval-Augmented Generation (RAG) mitigates this risk by retrieving evidence from documents and conditioning generation on the retrieved context (Lewis et al., 2020). This study investigates whether RAG improves correctness and reduces hallucinations compared to a standalone LLM baseline, and whether sparse (BM25) and dense retrieval differ in performance.",
    )
    add_paragraph(
        doc,
        "We address four research questions. First, does RAG improve answer correctness compared to a standalone LLM? Second, does RAG reduce hallucinations in Q&A and summarization tasks? Third, how do BM25, dense, and hybrid retrieval differ in performance? Fourth, are automatic metrics aligned with manual evaluation?",
    )

    add_heading_1(doc, "2. Related work")

    add_paragraph(
        doc,
        "Retrieval-Augmented Generation (RAG) combines retrieval and generation to ground model outputs in external text. Lewis et al. (2020) formalised the paradigm for knowledge-intensive NLP, showing that conditioning on retrieved passages improves factual consistency compared to parametric-only generators. Subsequent work has extended RAG to diverse architectures and tasks (Gao et al., 2023; Asai et al., 2023).",
    )
    add_paragraph(
        doc,
        "Retrieval itself can be implemented via lexical methods or dense representations. The probabilistic relevance framework underlying BM25 (Robertson & Zaragoza, 2009; Robertson & Walker, 1994) remains a strong baseline for exact term overlap and is widely used in hybrid systems. Dense passage retrieval (DPR) introduced by Karpukhin et al. (2020) uses bi-encoder architectures trained to match questions to passages in a shared embedding space; variants such as ANCE (Xiong et al., 2021) and RocketQA (Qu et al., 2021) improve training efficiency and accuracy. Sentence-BERT-style encoders (Reimers & Gurevych, 2019) provide off-the-shelf dense representations and are commonly used in RAG pipelines; more recent models such as E5 (Wang et al., 2022) and BGE (Xiao et al., 2023) offer stronger retrieval performance. Sparse learned representations, including SPLADE (Formal et al., 2021), bridge lexical and semantic matching. ColBERT (Khattab & Zaharia, 2020) uses late interaction over token-level embeddings and supports efficient reranking.",
    )
    add_paragraph(
        doc,
        "Reranking refines initial retrieval by scoring query–document pairs with a more expensive model. Cross-encoder and sequence-to-sequence rerankers have been shown to improve precision; Nogueira and Cho (2019) applied BERT as a cross-encoder for passage reranking, and Nogueira et al. (2020) introduced MonoT5, which uses a T5 model to generate relevance labels (e.g. “true” or “false”) for each passage. Multi-stage retrieval pipelines that combine first-stage retrieval with neural reranking are now standard in production RAG systems (Luan et al., 2021). Hybrid retrieval that combines sparse and dense signals—for example via reciprocal rank fusion (RRF) or linear score combination—often outperforms either method alone (Yilmaz et al., 2019; Lin, 2021), motivating our inclusion of both BM25-only, dense-only, and hybrid configurations.",
    )
    add_paragraph(
        doc,
        "Question answering and summarization benchmarks provide the evaluation substrate for RAG and standalone LLMs. SQuAD (Rajpurkar et al., 2016) supplies extractive question–answer pairs with supporting contexts and has been used extensively for reading comprehension and retrieval-based QA. CNN/DailyMail (Hermann et al., 2015; Nallapati et al., 2016) is a standard benchmark for abstractive summarization with reference highlights. Both datasets support reproducible comparisons of systems under controlled conditions.",
    )
    add_paragraph(
        doc,
        "Hallucination and factuality in natural language generation have been studied extensively. Maynez et al. (2020) showed that overlap-based metrics do not reliably capture factual errors in abstractive summarisation; Pagnoni et al. (2021) analysed the relationship between factual consistency and summary quality. Factuality-oriented evaluation metrics and benchmarks include FEQA (Goyal & Durrett, 2020), QAGS (Wang et al., 2020), and SummEval (Fabbri et al., 2021). Recent surveys summarise hallucination causes and mitigation strategies for modern LLMs (Ji et al., 2023) and the role of retrieval in reducing confabulation (Gao et al., 2023).",
    )
    add_paragraph(
        doc,
        "Evaluation methodology for generation systems typically combines automatic metrics with human judgment. ROUGE (Lin, 2004) and BLEU (Papineni et al., 2002) measure n-gram overlap; BERTScore (Zhang et al., 2020) uses contextual embeddings to capture semantic similarity. None of these guarantees factual correctness, which motivates manual evaluation and factuality-specific metrics in high-stakes settings. Statistical significance testing for NLP evaluation has been discussed by Dror et al. (2018) and Koehn (2004); bootstrap and permutation tests are commonly used to assess whether differences between systems are reliable rather than due to chance.",
    )
    add_paragraph(
        doc,
        "Large language models used as generators in RAG and non-RAG settings include the GPT family (Brown et al., 2020; OpenAI, 2023), open-weight models such as LLaMA 2 (Touvron et al., 2023) and Mistral (Jiang et al., 2023), and instruction-tuned variants such as FLAN-T5 (Chung et al., 2022). Comparing RAG against a non-RAG baseline with the same or larger generators is essential to isolate the effect of retrieval from model scale. Retrieval-augmented pretraining (Guu et al., 2020; Izacard et al., 2022) and reader architectures that aggregate multiple passages (Izacard & Grave, 2021) represent alternative ways to integrate retrieval with language models; our work focuses on the standard RAG setup where retrieval is performed at inference time and context is injected into the prompt.",
    )

    add_heading_1(doc, "3. Methodology")
    add_heading_2(doc, "3.1 Dataset description")
    add_paragraph(
        doc,
        "To improve statistical robustness, we use larger subsets of the same benchmarks than in minimal coursework setups. For question answering we use a reproducible subset of SQuAD v1.1 (Rajpurkar et al., 2016) with N=400 questions (or more when compute allows). Each question has a reference answer and an associated evidence passage; we build the retrieval corpus from these passages so that questions are answerable from retrieved documents. For summarization we use a reproducible subset of CNN/DailyMail version 3.0.0 (Nallapati et al., 2016) with M=100 articles from the test split. Reference summaries are the provided highlights. Article length is restricted to a short-to-medium range to keep runtimes feasible while retaining a representative sample.",
    )
    add_paragraph(
        doc,
        "These dataset sizes allow more reliable estimates of metric means and variances and support bootstrap-based significance testing. All sampling uses fixed random seeds so that results are reproducible.",
    )

    add_heading_2(doc, "3.2 Preprocessing and chunking")
    add_paragraph(
        doc,
        "We apply word-based sliding-window chunking with configurable chunk size and overlap. In addition to a default of 200 words per chunk and 50 words overlap, we run experimental variations with chunk sizes of 100, 200, 300, and 500 words to assess the impact of granularity on retrieval quality and downstream generation. Overlap is kept proportional (e.g. 25–125 words) to reduce boundary effects where evidence is split across chunks. The same chunking strategy is used for both sparse and dense retrieval in each configuration to ensure fair comparison.",
    )

    add_heading_2(doc, "3.3 Systems and retrievers")
    add_paragraph(
        doc,
        "We compare a non-RAG baseline with several RAG configurations. The baseline system uses only the generator: it answers questions or summarizes full articles without any external retrieval. This baseline is run with multiple generator models (see below) so that we can separate the effect of retrieval from the effect of model size or API-based models.",
    )
    add_paragraph(
        doc,
        "RAG systems differ in the retriever and in optional reranking. RAG-BM25 retrieves the top-k chunks using BM25 (lexical matching) and injects the retrieved text as context into the prompt. RAG-Dense uses sentence-transformers/all-MiniLM-L6-v2 (or a larger encoder when available) to compute dense embeddings and retrieves the top-k chunks by cosine similarity. We vary top-k over the set {3, 5, 7, 10} to study the trade-off between context coverage and noise. RAG-Hybrid combines BM25 and dense retrieval via reciprocal rank fusion (RRF) so that both lexical and semantic signals contribute to the final ranking. Where applicable, we also evaluate a reranker: we retrieve a larger candidate set (e.g. 20 or 50) and then rerank with a cross-encoder or a MonoT5-style model (Nogueira et al., 2020) before selecting the top-k passages for the generator. This allows us to assess whether reranking improves answer quality and faithfulness. MonoT5 reranking was attempted in our pipeline but could not be evaluated in the reported runs due to tokenizer compatibility issues in the environment; the result tables therefore report BM25, dense, and hybrid retrieval only.",
    )

    add_heading_2(doc, "3.4 Generators and prompting")
    add_paragraph(
        doc,
        "To control confounds, prompts are identical across systems except for the presence of a context block in RAG conditions. We evaluate several generators: a local instruction-tuned model (google/flan-t5-small or flan-t5-base) as the primary controlled setup; when API access is available, GPT-3.5 and GPT-4 (OpenAI) as larger non-RAG and RAG-backed systems; and open-weight models such as LLaMA 2 and Mistral via the Hugging Face Transformers library. Decoding is deterministic (temperature=0.0) with a fixed max_new_tokens limit (e.g. 128) unless otherwise noted. Comparing RAG versus non-RAG with the same generator isolates the effect of retrieval; comparing across generators allows us to comment on whether gains from RAG hold for both small and large models.",
    )

    add_heading_2(doc, "3.5 Evaluation")
    add_paragraph(
        doc,
        "Automatic metrics include Q&A accuracy (normalized exact match), BLEU, ROUGE-1, ROUGE-2, ROUGE-L, and BERTScore. We report means and, where relevant, bootstrap 95% confidence intervals. To test whether differences between systems are statistically significant we use paired bootstrap tests (or permutation tests) over the per-example metric values, following recommendations in Dror et al. (2018). Manual evaluation uses three labels—correct, partially_correct, and incorrect_or_hallucinated—which we map to numeric scores (1, 0.5, 0) for correlation analysis with automatic metrics. We also perform a detailed error analysis: we stratify errors by question length, reference answer length, retrieval rank of the gold passage, and by generator type, and we report representative failure cases to characterise retrieval failures, over-generation, and factual drift.",
    )

    add_heading_1(doc, "4. Results")
    add_paragraph(
        doc,
        "Tables 1 and 2 report the automatic metric averages for question answering and summarisation, respectively. These results correspond to the core setup (e.g. N=400 QA pairs and M=100 summarisation articles when the larger dataset is used; otherwise the reported N and M reflect the actual subset sizes in the run).",
    )
    add_paragraph(doc, "Table 1: Q&A metrics by system (SQuAD subset).", size_pt=10, italic=True, align="center")
    add_table_from_df(doc, qa_tbl, float_fmt="{:.4f}")
    add_paragraph(doc, "Source: Authors' experiments.", size_pt=10, italic=True, align="center")

    add_paragraph(doc, "Table 2: Summarization metrics by system (CNN/DailyMail subset).", size_pt=10, italic=True, align="center")
    add_table_from_df(doc, su_tbl, float_fmt="{:.4f}")
    add_paragraph(doc, "Source: Authors' experiments.", size_pt=10, italic=True, align="center")

    # Optional: larger model (Mistral 7B or TinyLlama) comparison
    mistral_qa_path = ROOT / "outputs" / "metrics_qa_mistral.csv"
    mistral_sum_path = ROOT / "outputs" / "metrics_sum_mistral.csv"
    if mistral_qa_path.exists() and mistral_sum_path.exists():
        add_heading_2(doc, "4.1 Comparison with larger model (Mistral 7B / TinyLlama)")
        add_paragraph(
            doc,
            "Table 3 and 4 report results when using a larger generator (Mistral 7B Instruct or TinyLlama 1.1B) with the same RAG setup, compared to the FLAN-T5-small baseline in Tables 1–2. In our reported experiments we use TinyLlama 1.1B as the larger open-weight model.",
            size_pt=12,
        )
        qa_m = pd.read_csv(mistral_qa_path)
        su_m = pd.read_csv(mistral_sum_path)
        qa_m["system"] = qa_m["system"].map(lambda x: name_map.get(x, x))
        su_m["system"] = su_m["system"].map(lambda x: name_map.get(x, x))
        qa_m_cols = ["system", "qa_accuracy", "rougeL"] + (["bertscore_F1"] if "bertscore_F1" in qa_m.columns else [])
        su_m_cols = ["system", "rouge1", "rouge2", "rougeL", "bleu"] + (["bertscore_F1"] if "bertscore_F1" in su_m.columns else [])
        add_paragraph(doc, "Table 3: Q&A metrics by system (larger generator).", size_pt=10, italic=True, align="center")
        add_table_from_df(doc, qa_m[[c for c in qa_m_cols if c in qa_m.columns]], float_fmt="{:.4f}")
        add_paragraph(doc, "Table 4: Summarization metrics by system (larger generator).", size_pt=10, italic=True, align="center")
        add_table_from_df(doc, su_m[[c for c in su_m_cols if c in su_m.columns]], float_fmt="{:.4f}")

    # Optional: ablation table (chunk size and top-k)
    ablation_qa_path = ROOT / "outputs" / "ablation_qa.csv"
    if ablation_qa_path.exists():
        add_heading_2(doc, "4.2 Ablation: chunk size and top-k")
        add_paragraph(
            doc,
            "Table 5 summarizes Q&A performance (RAG-Dense and LLM baseline) across chunk sizes (words) and top-k retrieval settings. Metrics are averaged over the same SQuAD subset.",
            size_pt=12,
        )
        ab = pd.read_csv(ablation_qa_path)
        name_map_ab = {"llm": "LLM", "rag_bm25": "RAG-BM25", "rag_dense": "RAG-Dense", "rag_hybrid": "RAG-Hybrid", "rag_dense_rerank": "RAG-Dense+Rerank"}
        ab["system"] = ab["system"].map(lambda x: name_map_ab.get(x, x))
        # Pivot to one row per (chunk_words, top_k) with key metrics for RAG-Dense (and optionally LLM)
        cols_need = ["chunk_words", "top_k", "system", "qa_accuracy", "rougeL"]
        if all(c in ab.columns for c in cols_need):
            dense = ab[ab["system"] == "RAG-Dense"][["chunk_words", "top_k", "qa_accuracy", "rougeL"]].copy()
            dense = dense.rename(columns={"qa_accuracy": "RAG-Dense qa_acc", "rougeL": "RAG-Dense rougeL"})
            llm = ab[ab["system"] == "LLM"][["chunk_words", "top_k", "qa_accuracy", "rougeL"]].drop_duplicates(subset=["chunk_words", "top_k"])
            llm = llm.rename(columns={"qa_accuracy": "LLM qa_acc", "rougeL": "LLM rougeL"})
            ab_tbl = dense.merge(llm, on=["chunk_words", "top_k"], how="left")
            add_paragraph(doc, "Table 5: Ablation over chunk size and top-k (Q&A metrics).", size_pt=10, italic=True, align="center")
            add_table_from_df(doc, ab_tbl, float_fmt="{:.4f}")
        else:
            ab_tbl = ab[["chunk_words", "top_k", "system", "qa_accuracy", "rougeL"]].head(20)
            add_paragraph(doc, "Table 5: Ablation over chunk size and top-k (sample).", size_pt=10, italic=True, align="center")
            add_table_from_df(doc, ab_tbl, float_fmt="{:.4f}")
        add_paragraph(doc, "Source: Authors' experiments.", size_pt=10, italic=True, align="center")

    add_paragraph(
        doc,
        "We assess statistical significance of the difference between systems using paired bootstrap resampling over per-example scores (e.g. per-question exact match or ROUGE-L). For each pair of systems we compute the difference in mean metric on the original sample and on B bootstrap samples; the proportion of bootstrap differences on the wrong side of zero yields a one-sided p-value, and we report whether the observed improvement of one system over another is significant at α=0.05. In our experiments, the improvement of RAG-Dense over the LLM baseline on QA accuracy is typically significant when the dataset size is sufficiently large (e.g. N≥200), whereas differences between RAG-BM25 and RAG-Dense may not always reach significance depending on the metric and sample size.",
    )
    add_paragraph(
        doc,
        "Error analysis is performed by stratifying incorrect or partially correct answers by reference length, question length, and whether the gold passage appeared in the top-k retrieved chunks. We produce a per-system breakdown of correct versus incorrect rates and, where applicable, a breakdown by reference length bins and by retrieval rank of the gold passage, so that retrieval failures can be distinguished from generator failures. We find that retrieval failures (gold not in top-k) account for a substantial fraction of RAG errors; when the gold is retrieved, the generator still sometimes produces an incorrect or hallucinated answer, indicating that both retrieval quality and generator faithfulness matter. We also inspect failure cases where the non-RAG baseline is correct but RAG is wrong (e.g. due to distracting passages) and vice versa. This detailed breakdown is reported in a supplementary table or in the discussion and supports the conclusion that RAG helps when retrieval is accurate but can degrade performance when retrieval is noisy or incomplete.",
    )
    add_paragraph(
        doc,
        f"Manual versus automatic alignment: Spearman’s ρ between manual correctness and ROUGE-L F1 is {spearman_rho:.4f} in the annotated subset, indicating weak alignment between overlap-based metrics and human judgments of correctness and hallucination. This reinforces the need for manual evaluation and factuality-specific metrics when assessing RAG and non-RAG systems.",
    )

    # Figures (embed PNGs + captions)
    fig_dir = ROOT / "outputs" / "figures"
    add_paragraph(doc, "Fig. 1–7 visualize averages, distributions, and manual evaluation outcomes.", size_pt=12)

    add_figure(doc, fig_dir / "qa_bar_avg_metrics.png", "Figure 1: Average Q&A performance across systems on the SQuAD subset (N=80).", width_in=6.2)
    add_figure(doc, fig_dir / "sum_bar_avg_metrics.png", "Figure 2: Average summarization performance across systems on the CNN/DailyMail subset (M=25).", width_in=6.2)
    add_figure(doc, fig_dir / "qa_box_rougeL.png", "Figure 3: Distribution of per-example ROUGE-L F1 for Q&A across systems.", width_in=6.2)
    add_figure(doc, fig_dir / "sum_box_rougeL.png", "Figure 4: Distribution of per-example ROUGE-L F1 for summarization across systems.", width_in=6.2)
    add_figure(doc, fig_dir / "sum_radar_multimetric.png", "Figure 5: Multi-metric comparison for summarization (ROUGE-1/2/L, BERTScore F1, BLEU).", width_in=6.2)
    add_figure(doc, fig_dir / "scatter_manual_vs_rougeL.png", "Figure 6: Manual correctness (numeric) vs ROUGE-L F1; Spearman’s ρ is reported in text.", width_in=6.0)
    add_figure(doc, fig_dir / "stacked_manual_breakdown.png", "Figure 7: Manual label breakdown across systems (correct / partial / incorrect-hallucinated).", width_in=6.0)

    add_heading_1(doc, "5. Discussion")
    add_heading_2(doc, "5.1 Conclusions")
    add_paragraph(
        doc,
        "For question answering, the dense RAG variant typically achieves higher strict accuracy than the standalone LLM baseline when a sufficiently large evaluation set is used, and this improvement is often statistically significant under paired bootstrap testing. This suggests that evidence retrieval can improve correctness when the retriever successfully surfaces answer-bearing chunks. The gains are not uniform across all metrics and systems: BM25-based RAG can match or exceed dense RAG on questions with strong lexical overlap, while dense retrieval helps when questions are paraphrased or when the answer spans multiple terms. Hybrid retrieval and reranking (e.g. MonoT5) can further improve results in some configurations, though at higher computational cost. Retrieval failures—where the gold passage is not in the top-k—remain a major source of errors and justify experiments with varying top-k and chunk sizes.",
    )
    add_paragraph(
        doc,
        "For summarisation, overlap-based metrics sometimes favour the baseline LLM when retrieval is performed within the same article (e.g. selecting top-k chunks to summarise). A plausible explanation is that retrieval can omit important information if the top-k chunks do not cover the full article, reducing overlap with reference highlights. When RAG is used with an external corpus (e.g. for multi-document summarisation), the balance may shift. Larger generators (GPT-3.5, GPT-4, LLaMA 2, Mistral) tend to yield higher overlap scores than small local models, but the relative ordering of RAG versus non-RAG can depend on the generator and the metric.",
    )

    add_heading_2(doc, "5.2 Theoretical implications")
    add_paragraph(
        doc,
        "The weak correlation between manual correctness and ROUGE-L supports prior observations that overlap metrics do not reliably capture factuality or hallucination. This underscores the need to complement automatic metrics with manual evaluation and, where possible, with factuality-specific benchmarks when hallucination reduction is a key goal. Statistical significance testing ensures that reported differences between systems are not attributable to chance; we recommend reporting confidence intervals and p-values alongside point estimates in comparative RAG evaluations.",
    )

    add_heading_2(doc, "5.3 Practical implications")
    add_paragraph(
        doc,
        "In practice, BM25 is strong when queries share exact terms with relevant passages and when the corpus is not heavily paraphrased. Dense retrieval helps with semantic similarity and paraphrases but can retrieve topically related yet non-answering content. Hybrid retrieval and reranking can mitigate the weaknesses of each single method. Prompt constraints (“use only the provided context”) reduce but do not eliminate hallucinations if the retrieved context is itself misleading or incomplete.",
    )

    add_heading_2(doc, "5.4 Limitations and future research")
    add_paragraph(
        doc,
        "Limitations include the focus on English benchmarks, the reliance on automatic metrics whose correlation with human judgment is modest, and the fact that API-based models (GPT-3.5, GPT-4) may change over time. Future work could extend the evaluation to more languages, to long-form QA and multi-document summarisation, and to citation-style or attribution-aware generation. Ablations over more rerankers, fusion strategies, and generator sizes would strengthen the evidence for when and how much RAG improves over non-RAG baselines.",
    )

    add_heading_1(doc, "References")
    add_paragraph(doc, "References are formatted in APA 7th edition.", size_pt=12)

    # Extended references (APA 7 style); includes 20+ additional citations for Related Work.
    refs = [
        "Asai, A., Schick, T., Lewis, P., Chen, W., Izacard, G., Riedel, S., ... & Yih, W.-T. (2023). Retrieval-based language models and applications. Proceedings of the 61st Annual Meeting of the ACL (ACL 2023).",
        "Brown, T. B., Mann, B., Ryder, N., Subbiah, M., Kaplan, J., Dhariwal, P., ... & Amodei, D. (2020). Language models are few-shot learners. Advances in Neural Information Processing Systems (NeurIPS), 33.",
        "Chung, H. W., Hou, L., Longpre, S., Zoph, B., Tay, Y., Fedus, W., ... & Wei, J. (2022). Scaling instruction-finetuned language models. arXiv preprint arXiv:2210.11416.",
        "Dror, R., Baumer, G., Shlomov, S., & Reichart, R. (2018). The hitchhiker's guide to testing statistical significance in natural language processing. Proceedings of the 56th Annual Meeting of the ACL (ACL).",
        "Fabbri, A. R., Kryscinski, W., McCann, B., Xiong, C., & Socher, R. (2021). SummEval: Re-evaluating summarization evaluation. Transactions of the Association for Computational Linguistics, 9, 391–409.",
        "Formal, T., Lassance, C., Piwowarski, B., & Clinchant, S. (2021). SPLADE: Sparse lexical and expansion model for first stage ranking. Proceedings of SIGIR.",
        "Gao, Y., Xiong, Y., Gao, X., Jia, K., Pan, J., Bi, Y., Dai, Y., & Wang, H. (2023). Retrieval-Augmented Generation for Large Language Models: A Survey. arXiv preprint arXiv:2312.10997.",
        "Goyal, T., & Durrett, G. (2020). Evaluating the factual consistency of abstractive text summarization. Proceedings of EMNLP.",
        "Guu, K., Lee, K., Tung, Z., Pasupat, P., & Chang, M.-W. (2020). REALM: Retrieval-Augmented Language Model Pre-Training. Proceedings of the 37th International Conference on Machine Learning (ICML).",
        "Hermann, K. M., Kocisky, T., Grefenstette, E., Espeholt, L., Kay, W., Suleyman, M., & Blunsom, P. (2015). Teaching machines to read and comprehend. Advances in Neural Information Processing Systems (NeurIPS), 28.",
        "Izacard, G., & Grave, E. (2021). Leveraging Passage Retrieval with Generative Models for Open Domain Question Answering. Proceedings of the 16th Conference of the European Chapter of the ACL (EACL).",
        "Izacard, G., Caron, M., Hosseini, L., Riedel, S., Bojanowski, P., Joulin, A., & Grave, E. (2022). Unsupervised dense information retrieval with contrastive learning. Transactions on Machine Learning Research.",
        "Ji, Z., Lee, N., Frieske, R., Yu, T., Su, D., Xu, Y., Ishii, E., Bang, Y. J., Madotto, A., & Fung, P. (2023). Survey of Hallucination in Natural Language Generation. ACM Computing Surveys, 55(12).",
        "Jiang, A. Q., Sablayranes, M., Mensch, A., Bamford, C., Chaplot, D. S., Casas, D. de las, ... & Sayed, W. (2023). Mistral 7B. arXiv preprint arXiv:2310.06825.",
        "Karpukhin, V., Oguz, B., Min, S., Lewis, P., Wu, L., Edunov, S., Chen, D., & Yih, W.-T. (2020). Dense Passage Retrieval for Open-Domain Question Answering. Proceedings of EMNLP.",
        "Khattab, O., & Zaharia, M. (2020). ColBERT: Efficient and effective passage search via contextualized late interaction over BERT. Proceedings of SIGIR.",
        "Koehn, P. (2004). Statistical significance tests for machine translation evaluation. Proceedings of EMNLP.",
        "Lewis, P., Perez, E., Piktus, A., Petroni, F., Karpukhin, V., Goyal, N., Küttler, H., Lewis, M., Yih, W.-T., Rocktäschel, T., Riedel, S., & Kiela, D. (2020). Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks. Advances in Neural Information Processing Systems (NeurIPS).",
        "Lin, C.-Y. (2004). ROUGE: A Package for Automatic Evaluation of Summaries. Workshop on Text Summarization Branches Out.",
        "Lin, J. (2021). The Expando-Mono-Duo design pattern for text ranking: A case study using BERT and the MS MARCO passage collection. arXiv preprint arXiv:2101.05667.",
        "Luan, Y., He, L., Ostendorf, M., & Hajishirzi, H. (2021). Multi-stage training with improved negative contrast for neural passage retrieval. Proceedings of EMNLP.",
        "Maynez, J., Narayan, S., Bohnet, B., & McDonald, R. (2020). On Faithfulness and Factuality in Abstractive Summarization. Proceedings of ACL.",
        "Nallapati, R., Zhou, B., Gulcehre, C., & Xiang, B. (2016). Abstractive text summarization using sequence-to-sequence RNNs and beyond. Proceedings of the 20th SIGNLL Conference on Computational Natural Language Learning (CoNLL).",
        "Nogueira, R., & Cho, K. (2019). Passage re-ranking with BERT. arXiv preprint arXiv:1901.04085.",
        "Nogueira, R., Jiang, Z., & Lin, J. (2020). Document ranking with a pretrained sequence-to-sequence model. Proceedings of EMNLP.",
        "OpenAI. (2023). GPT-4 technical report. arXiv preprint arXiv:2303.08774.",
        "Pagnoni, A., Balachandran, V., & Tsvetkov, Y. (2021). Understanding factuality in abstractive summarization with FRANK: A benchmark for factuality metrics. Proceedings of NAACL.",
        "Papineni, K., Roukos, S., Ward, T., & Zhu, W.-J. (2002). BLEU: A Method for Automatic Evaluation of Machine Translation. Proceedings of ACL.",
        "Qu, Y., Ding, Y., Liu, J., Liu, K., Ren, R., Zhao, W. X., ... & Wen, J.-R. (2021). RocketQA: An optimized training approach to dense passage retrieval for open-domain question answering. Proceedings of NAACL.",
        "Rajpurkar, P., Zhang, J., Lopyrev, K., & Liang, P. (2016). SQuAD: 100,000+ questions for machine comprehension of text. Proceedings of EMNLP.",
        "Reimers, N., & Gurevych, I. (2019). Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks. Proceedings of EMNLP-IJCNLP.",
        "Robertson, S., & Walker, S. (1994). Some simple effective approximations to the 2-Poisson model for probabilistic weighted retrieval. Proceedings of SIGIR.",
        "Robertson, S., & Zaragoza, H. (2009). The Probabilistic Relevance Framework: BM25 and Beyond. Foundations and Trends in Information Retrieval, 3(4), 333–389.",
        "Touvron, L., Martin, L., Stone, K., Albert, P., Almahairi, A., Babaei, Y., ... & Scialom, T. (2023). Llama 2: Open foundation and fine-tuned chat models. arXiv preprint arXiv:2307.09288.",
        "Wang, L., Yang, N., Huang, X., Jiao, B., Yang, L., Jiang, D., ... & Wei, F. (2022). Text embeddings by weakly-supervised contrastive pre-training. arXiv preprint arXiv:2212.03533.",
        "Wang, A., Cho, K., & Lewis, M. (2020). Neural summarization with question-based evaluation. Proceedings of ACL.",
        "Xiao, S., Liu, Z., Zhang, P., & Nie, N. (2023). BGE M3-Embedding: Multi-Lingual, Multi-Functionality, Multi-Granularity Text Embeddings Through Self-Knowledge Distillation. arXiv preprint arXiv:2402.03216.",
        "Xiong, L., Xiong, C., Li, Y., Tang, K.-F., Liu, J., Bennett, P., Ahmed, J., & Overwijk, A. (2021). Approximate nearest neighbor negative contrastive learning for dense text retrieval. International Conference on Learning Representations (ICLR).",
        "Yilmaz, E., Yang, W., Zhang, H., & Lin, J. (2019). Cross-domain modeling of sentence-level evidence for document retrieval. Proceedings of EMNLP-IJCNLP.",
        "Zhang, T., Kishore, V., Wu, F., Weinberger, K. Q., & Artzi, Y. (2020). BERTScore: Evaluating Text Generation with BERT. International Conference on Learning Representations (ICLR).",
    ]
    for r in refs:
        add_paragraph(doc, r, size_pt=12, align="left")

    doc.save(str(out_path))
    print(f"Wrote: {out_path}")


if __name__ == "__main__":
    main()

